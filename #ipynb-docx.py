#%%
import os
import shutil
import docx
import datetime
from docx.oxml.ns import qn
from tkinter import messagebox

s = datetime.datetime.now()  #当前时间
date = s.strftime("%Y%m%d")
#-------------------------------------------------------------------
source_path = os.path.abspath('C:/Users/Administrator')  #源地址
target_path = os.path.abspath('D:/ipynb')  #目标地址

if not os.path.exists(target_path):  # 如不存在目标地址
    os.makedirs(target_path)  #建立

for file in os.listdir(target_path):  
    os.remove(target_path + '/' + file)

for file in os.listdir(source_path):  # 遍历文件夹

    try:
        if file.split('.')[1] == str('ipynb'):
            src_file = os.path.join(source_path, file)
            shutil.copy(src_file, target_path)
            os.remove(src_file)
            
    except IndexError:
        pass


os.chdir(target_path)
command = 'jupyter nbconvert --to script *.ipynb '
os.system(command)

for file in os.listdir(target_path):  # 遍历文件夹

    try:
        if file.split('.')[1] == str('ipynb'):
            os.remove(target_path + '/' + file)

        elif file.split('.')[1] == str('py'):

            o = target_path + '/' + file
            n = target_path + '/' + os.path.splitext(file)[0] + '.txt'

            os.rename(o, n)            


    except IndexError:
        pass




document = docx.Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

for file in os.listdir(target_path):  # 遍历文件夹
    text = open(target_path+'/' + file, encoding='utf-8')  #依次打开
    rl = text.readlines()  #读取所有行
    paragraph = document.add_paragraph(os.path.splitext(file)[0],'Heading 1')  #提取文件名
    paragraph = document.add_paragraph(''.join(rl[0:len(rl)]), 'Body Text')
    text.close()
document.save("D:/OneDrive/data/ipynb" + date + ".docx")

#-------------------------------------------------------------------

#os.chdir('c:/in')
#shutil.rmtree(target_path)

e = datetime.datetime.now()  #当前时间                                                                             

messagebox.showinfo("提示", "用时:" + str(e - s) + '秒')

# %%
