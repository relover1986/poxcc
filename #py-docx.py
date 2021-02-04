#%%
import os
import shutil
import docx
import datetime
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt # 导入单位换算函数



  
s = datetime.datetime.now()  #当前时间
date=s.strftime("%Y%m%d")
#-------------------------------------------------------------------     

time=s.strftime("%Y%m%d%H%M%S")

source_path = os.path.abspath('D:/OneDrive/python')#源地址
target_path = os.path.abspath('D:/time')#目标地址
if not os.path.exists(target_path):# 如不存在目标地址
    os.makedirs(target_path)#建立
    
for file in os.listdir(source_path):# 遍历文件夹   
    
    try:
        if file.split('.')[1]==str('py'):
            src_file = os.path.join(source_path, file)
            shutil.copy(src_file, target_path)
            #print(file)
    except IndexError:       
        pass
path = target_path+"/"

for i in os.listdir(path):
    o = path + i
    n = path + os.path.splitext(i)[0]+'.txt'
    
    os.rename(o,n)
    


document = docx.Document()
section = document.sections[0] # 获取section对象
section.top_margin = Cm(1.27)
section.bottom_margin = Cm(1.27)
section.left_margin = Cm(1.27)
section.right_margin = Cm(1.27)
document.styles['Normal'].font.name = u'楷体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
for file in os.listdir(path):# 遍历文件夹 
    if file.endswith('.txt'):
        text=open(path+file,encoding='utf-8')#依次打开
        rl=text.readlines()#读取所有行
        paragraph = document.add_paragraph(os.path.splitext(file)[0],'Heading 1')#提取文件名
        paragraph = document.add_paragraph(''.join(rl[0:len(rl)]),'Body Text')
        text.close()
document.save("D:/OneDrive/data/代码-"+date+".docx")


shutil.rmtree(path)

#------------------------------------------------------------------- 
    
e = datetime.datetime.now()  #当前时间

print("用时:" + str(e - s) + '秒')

# %%
